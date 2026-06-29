import io
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pdf_service import EvaluatorPDF

def export_to_word(questionnaire):
    """Exporte le questionnaire au format Word (docx) en respectant la charte visuelle."""
    doc = docx.Document()
    
    # Configuration des styles par défaut
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # En-tête / Titre du questionnaire
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(questionnaire.title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(99, 102, 241)  # Accent Indigo #6366f1
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if questionnaire.description:
        desc_p = doc.add_paragraph()
        desc_p.add_run(questionnaire.description).italic = True
        desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    doc.add_paragraph()  # Espaceur
    
    # Parcours des blocs du questionnaire
    for b in questionnaire.blocks:
        if b.block_type == 'title':
            # Déjà traité
            pass
        elif b.block_type == 'section':
            section_title = b.content.get('title', 'Section sans titre')
            p = doc.add_paragraph()
            run = p.add_run(section_title)
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(124, 58, 237)  # Violet
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
        elif b.block_type == 'text':
            text_content = b.content.get('text', '')
            p = doc.add_paragraph(text_content)
            p.paragraph_format.space_after = Pt(10)
        elif b.block_type == 'question':
            label = b.content.get('label', '')
            q_type = b.content.get('question_type', 'text')
            is_req = b.content.get('is_required', False)
            choices = b.content.get('choices', [])
            help_txt = b.content.get('help_text', '')
            
            p = doc.add_paragraph()
            run = p.add_run(f"❓ {label}")
            run.bold = True
            if is_req:
                p.add_run(" *").font.color.rgb = RGBColor(239, 68, 68)  # Étoile rouge si obligatoire
                
            if help_txt:
                hp = doc.add_paragraph()
                hrun = hp.add_run(f"Note : {help_txt}")
                hrun.italic = True
                hrun.font.size = Pt(9.5)
                hrun.font.color.rgb = RGBColor(100, 100, 100)
                hp.paragraph_format.left_indent = Inches(0.25)
                
            if q_type in ['select', 'radio', 'checkbox']:
                for c in choices:
                    cp = doc.add_paragraph(f"  [ ] {c}")
                    cp.paragraph_format.left_indent = Inches(0.4)
            else:
                cp = doc.add_paragraph("  __________________________________________________")
                cp.paragraph_format.left_indent = Inches(0.4)
                
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif b.block_type == 'signature':
            p = doc.add_paragraph()
            p.add_run("✍️ Signature :").bold = True
            cp = doc.add_paragraph("\n  ___________________________\n")
            cp.paragraph_format.left_indent = Inches(0.4)
        elif b.block_type == 'gps':
            p = doc.add_paragraph()
            p.add_run("📍 Coordonnées GPS : ________________________").bold = True
            p.paragraph_format.space_after = Pt(10)
        elif b.block_type == 'photo':
            p = doc.add_paragraph()
            p.add_run("📷 Capture Photo : [ Emplacement image ]").bold = True
            p.paragraph_format.space_after = Pt(10)
        elif b.block_type == 'file':
            p = doc.add_paragraph()
            p.add_run("📁 Pièce jointe requise : ____________________").bold = True
            p.paragraph_format.space_after = Pt(10)
        elif b.block_type == 'checkbox':
            label = b.content.get('label', 'Checklist')
            opts = b.content.get('options', [])
            p = doc.add_paragraph()
            p.add_run(f"☑️ {label}").bold = True
            for o in opts:
                cp = doc.add_paragraph(f"  [ ] {o}")
                cp.paragraph_format.left_indent = Inches(0.4)
        elif b.block_type == 'comment':
            p = doc.add_paragraph()
            p.add_run("💬 Commentaires libres :").bold = True
            doc.add_paragraph("\n\n\n").paragraph_format.left_indent = Inches(0.4)
            
    # Sauvegarde dans un flux mémoire
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream.getvalue()

def export_to_excel(questionnaire):
    """Exporte le questionnaire au format Excel (xlsx) en listant la structure complète."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "S&E-CSB Questionnaire"
    
    headers = ["Index", "Type de Bloc", "Contenu / Question", "Options / Choix", "Obligatoire", "Aide à la saisie"]
    ws.append(headers)
    
    # Styles Excel
    header_fill = PatternFill(start_color="6366f1", end_color="6366f1", fill_type="solid")
    header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    for col_num, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        
    idx = 1
    for b in questionnaire.blocks:
        if b.block_type == 'title':
            ws.append([idx, "Titre", b.content.get('title', ''), b.content.get('description', ''), "", ""])
        elif b.block_type == 'section':
            ws.append([idx, "Section", b.content.get('title', ''), "", "", ""])
        elif b.block_type == 'text':
            ws.append([idx, "Texte descriptif", b.content.get('text', ''), "", "", ""])
        elif b.block_type == 'question':
            choices = ", ".join(b.content.get('choices', []))
            ws.append([
                idx,
                f"Question ({b.content.get('question_type', 'text')})",
                b.content.get('label', ''),
                choices,
                "Oui" if b.content.get('is_required', False) else "Non",
                b.content.get('help_text', '')
            ])
        elif b.block_type == 'checkbox':
            opts = ", ".join(b.content.get('options', []))
            ws.append([idx, "Checklist", b.content.get('label', ''), opts, "", ""])
        else:
            ws.append([idx, b.block_type.upper(), b.content.get('label', ''), "", "", b.content.get('help_text', '')])
        idx += 1
        
    # Redimensionnement automatique des colonnes
    for col in ws.columns:
        max_len = max(len(str(cell.value or '')) for cell in col)
        col_letter = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[col_letter].width = max(max_len + 3, 10)
        
    stream = io.BytesIO()
    wb.save(stream)
    stream.seek(0)
    return stream.getvalue()

def export_to_pdf(questionnaire):
    """Exporte le questionnaire en PDF stylisé via EvaluatorPDF."""
    pdf = EvaluatorPDF(title_text=questionnaire.title, project_name="S&E-CSB")
    pdf.add_page()
    
    # Titre du Questionnaire
    pdf.set_font('helvetica', 'B', 18)
    pdf.set_text_color(99, 102, 241)  # Accent Indigo
    pdf.multi_cell(0, 10, questionnaire.title, 0, 'C')
    pdf.ln(4)
    
    # Description
    if questionnaire.description:
        pdf.set_font('helvetica', 'I', 11)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 6, questionnaire.description, 0, 'C')
        pdf.ln(8)
        
    # Parcours des blocs
    for b in questionnaire.blocks:
        if b.block_type == 'section':
            pdf.set_font('helvetica', 'B', 13)
            pdf.set_fill_color(240, 243, 248)
            pdf.set_text_color(124, 58, 237)  # Violet
            pdf.cell(0, 10, f"  {b.content.get('title', 'Section')}", 0, 1, 'L', fill=True)
            pdf.ln(4)
        elif b.block_type == 'text':
            pdf.set_font('helvetica', '', 10)
            pdf.set_text_color(60, 60, 60)
            pdf.multi_cell(0, 5, b.content.get('text', ''))
            pdf.ln(4)
        elif b.block_type == 'question':
            label = b.content.get('label', '')
            q_type = b.content.get('question_type', 'text')
            is_req = b.content.get('is_required', False)
            choices = b.content.get('choices', [])
            help_txt = b.content.get('help_text', '')
            
            pdf.set_font('helvetica', 'B', 11)
            pdf.set_text_color(0, 0, 0)
            req_str = " *" if is_req else ""
            pdf.multi_cell(0, 6, f"Q: {label}{req_str}")
            
            if help_txt:
                pdf.set_font('helvetica', 'I', 9)
                pdf.set_text_color(120, 120, 120)
                pdf.multi_cell(0, 4.5, f"Note : {help_txt}")
                
            pdf.ln(2)
            
            # Affichage de la zone de réponse
            pdf.set_font('helvetica', '', 10)
            pdf.set_text_color(80, 80, 80)
            if q_type in ['select', 'radio', 'checkbox']:
                for c in choices:
                    pdf.cell(10)
                    pdf.cell(0, 6, f"[ ] {c}", 0, 1, 'L')
            else:
                pdf.cell(10)
                pdf.cell(0, 6, "__________________________________________________", 0, 1, 'L')
            pdf.ln(5)
        elif b.block_type == 'signature':
            pdf.set_font('helvetica', 'B', 11)
            pdf.cell(0, 6, "Signature :", 0, 1, 'L')
            pdf.ln(12)
            # Dessiner une ligne pour signer
            pdf.line(pdf.get_x() + 10, pdf.get_y(), pdf.get_x() + 80, pdf.get_y())
            pdf.ln(6)
        elif b.block_type == 'gps':
            pdf.set_font('helvetica', 'B', 11)
            pdf.cell(0, 6, "📍 Coordonnees GPS : ________________________", 0, 1, 'L')
            pdf.ln(4)
        elif b.block_type == 'photo':
            pdf.set_font('helvetica', 'B', 11)
            pdf.cell(0, 6, "📷 Capture Photo : [ Cadre photo ]", 0, 1, 'L')
            pdf.ln(4)
        elif b.block_type == 'checkbox':
            label = b.content.get('label', 'Checklist')
            opts = b.content.get('options', [])
            pdf.set_font('helvetica', 'B', 11)
            pdf.cell(0, 6, f"☑️ {label}", 0, 1, 'L')
            pdf.set_font('helvetica', '', 10)
            for o in opts:
                pdf.cell(10)
                pdf.cell(0, 6, f"[ ] {o}", 0, 1, 'L')
            pdf.ln(4)
            
    return pdf.output()
